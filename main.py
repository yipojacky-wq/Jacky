from io import BytesIO

from docx import Document


def markdown_to_docx_bytes(markdown: str, title: str) -> bytes:
    document = Document()
    document.add_heading(title or "可撰寫揭露書初稿", level=0)
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("> "):
            paragraph = document.add_paragraph(line[2:])
            if paragraph.runs:
                paragraph.runs[0].italic = True
        else:
            document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
